"""DOCX renderer for the neutral report representation."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from atlas.reports.schemas import ReportRepresentation


def render_docx(report: ReportRepresentation) -> bytes:
    document = Document()
    document.core_properties.title = report.title
    document.core_properties.subject = "ATLAS evidence-backed report"
    document.add_heading(report.title, level=0)
    document.add_paragraph(f"Source run: {report.source_run_id}")
    for section in report.sections:
        document.add_heading(section.title, level=1)
        for line in section.narrative.splitlines() or [section.narrative]:
            document.add_paragraph(line)
    document.add_heading("Evidence manifest", level=1)
    for citation in report.citations:
        document.add_paragraph(
            f"[{citation.citation_id}] {citation.url} — Original evidence: {citation.excerpt}"
        )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
